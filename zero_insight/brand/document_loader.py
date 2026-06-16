from __future__ import annotations

import zipfile
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path


@dataclass
class LoadedBrandDocument:
    path: Path
    text: str
    warnings: list[str] = field(default_factory=list)
    extracted_assets: list[Path] = field(default_factory=list)


class DocumentLoader:
    def load(self, path: Path, assets_dir: Path | None = None) -> LoadedBrandDocument:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._load_pdf(path, assets_dir)
        if suffix == ".docx":
            return self._load_docx(path, assets_dir)
        if suffix in {".png", ".jpg", ".jpeg", ".webp"}:
            return self._load_image(path, assets_dir)
        raise ValueError("Formato nao suportado. Use PDF, DOCX, PNG, JPG ou WEBP.")

    def _load_pdf(self, path: Path, assets_dir: Path | None) -> LoadedBrandDocument:
        try:
            import fitz
        except ImportError as exc:
            raise RuntimeError("PyMuPDF nao instalado. Instale a dependencia pymupdf.") from exc

        warnings: list[str] = []
        assets: list[Path] = []
        chunks: list[str] = []
        doc = fitz.open(path)
        for page_index, page in enumerate(doc, start=1):
            chunks.append(page.get_text("text"))
            if assets_dir:
                assets_dir.mkdir(parents=True, exist_ok=True)
                for image_index, image in enumerate(page.get_images(full=True), start=1):
                    xref = image[0]
                    extracted = doc.extract_image(xref)
                    ext = extracted.get("ext", "png")
                    image_path = assets_dir / f"extracted_image_{page_index:02d}_{image_index:02d}.{ext}"
                    image_path.write_bytes(extracted["image"])
                    assets.append(image_path)
        text = "\n".join(chunks).strip()
        if not text:
            warnings.append("PDF sem texto extraivel. OCR nao faz parte do MVP.")
        return LoadedBrandDocument(path=path, text=text, warnings=warnings, extracted_assets=assets)

    def _load_docx(self, path: Path, assets_dir: Path | None) -> LoadedBrandDocument:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx nao instalado. Instale a dependencia python-docx.") from exc

        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip()).strip()
        assets: list[Path] = []
        if assets_dir:
            assets_dir.mkdir(parents=True, exist_ok=True)
            with zipfile.ZipFile(path) as archive:
                for name in archive.namelist():
                    if name.startswith("word/media/"):
                        target = assets_dir / Path(name).name
                        target.write_bytes(archive.read(name))
                        assets.append(target)
        warnings = [] if text else ["DOCX sem texto extraivel."]
        return LoadedBrandDocument(path=path, text=text, warnings=warnings, extracted_assets=assets)

    def _load_image(self, path: Path, assets_dir: Path | None) -> LoadedBrandDocument:
        try:
            from PIL import Image
        except ImportError as exc:
            raise RuntimeError("Pillow nao instalado. Instale a dependencia pillow.") from exc

        assets: list[Path] = []
        if assets_dir:
            assets_dir.mkdir(parents=True, exist_ok=True)
            target = assets_dir / path.name
            if path.resolve() != target.resolve():
                target.write_bytes(path.read_bytes())
            assets.append(target)

        colors: list[str] = []
        try:
            with Image.open(path) as image:
                image = image.convert("RGB")
                image.thumbnail((96, 96))
                pixels = [
                    (r, g, b)
                    for r, g, b in image.getdata()
                    if not (r > 245 and g > 245 and b > 245) and not (r < 10 and g < 10 and b < 10)
                ]
                for (r, g, b), _ in Counter(pixels).most_common(6):
                    colors.append(f"#{r:02X}{g:02X}{b:02X}")
        except Exception as exc:
            return LoadedBrandDocument(
                path=path,
                text=f"Arquivo visual da marca: {path.stem}",
                warnings=[f"Nao foi possivel analisar cores da imagem: {exc}"],
                extracted_assets=assets,
            )

        text = "\n".join(
            [
                f"Arquivo visual da marca: {path.stem}",
                "Visual: imagem/logo importado para uso como asset da marca.",
                f"Cores identificadas: {', '.join(colors) if colors else 'nenhuma cor dominante identificada'}",
                "Logo: usar a imagem importada como referencia visual principal.",
            ]
        )
        return LoadedBrandDocument(path=path, text=text, warnings=[], extracted_assets=assets)
